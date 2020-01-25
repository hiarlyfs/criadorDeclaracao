from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dataAtual import dataAtual
from datetime import datetime
from aluno import *
import json

now = datetime.now()
arquivo = open('arquivo.json', 'r')
dados = json.loads(arquivo.readline())

def addTexto(str):
    run = paragrafo.add_run(str)
    font = run.font
    font.name = 'Arial Narrow'
    font.size = Pt(14)


def addTextoNegrito(str):
    run = paragrafo.add_run(str)
    font = run.font
    font.name = 'Arial Narrow'
    font.size = Pt(14)
    font.bold = True

def validarDataNascimento(str):
    data = str.split('/')
    if len(data) != 3:
        return False
    if len(data[0]) != 2 or len(data[1]) != 2 or len(data[2]) != 4:
        return False
    try:
        if 1 < int(data[0]) > 31:
            return False
        if 1 < int(data[1]) > 12:
            return False
        return True
    except ValueError:
        return False 



document = Document('.modelo.docx')
aluno = dados['nomeAluno'].upper()

while True:
    dataNascimento = dados['dataNascimento']
    if (validarDataNascimento(dataNascimento)): break
    else: 
        print('Data de Nascimento Inválida... Exemplo Correto: 01/01/2001')
        print()

nomePai = dados['nomePai'].upper()
nomeMae = dados['nomeMae'].upper()

cursou = dados['serieCursou']
cursara = proximaSerie(cursou)

paragrafo = document.add_paragraph()
addTexto("\tDeclaro para os devidos fins de comprovação que a(o) aluna(o) ")
addTextoNegrito(aluno  + " ")
addTexto("nascida(o) no dia ")
addTextoNegrito(dataNascimento + " ")
addTexto('filha(o) de  ')
addTextoNegrito(nomePai + " E " + nomeMae + " ")
addTexto('cursou e frequentou o ')
addTextoNegrito(cursou + " ")
addTexto('na Escola Municipal de Ensino Fundamental José Lopes no ano letivo de ' + str(now.year-1) +  ' e está apto(a) a cursar o ')
addTextoNegrito(cursara)
addTexto('.')
paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragrafo.paragraph_format.line_spacing = 2.0

document.add_paragraph()

paragrafo = document.add_paragraph()
addTexto(dataAtual())
paragrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT

document.add_paragraph()
document.add_paragraph()

paragrafo = document.add_paragraph()
addTextoNegrito('______________________________________')
paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

paragrafo = document.add_paragraph()
addTextoNegrito('ADMINISTRADOR(A) ESCOLAR')
paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER


document.save(aluno.replace(" ", "_") + ".docx")
